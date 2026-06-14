from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from generate_algorithm_design_doc import (
    add_code_block,
    add_footer,
    add_heading,
    add_list,
    add_para,
    add_table,
    repair_settings_zoom,
    set_cell_text,
    set_cell_width,
    set_full_grid_borders,
    set_run_font,
)


OUT = Path("一号线日周排产系统_数据设计说明书_已填写.docx")


def add_grid_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths is None
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        if widths:
            set_cell_width(cell, widths[i])
        set_cell_text(cell, h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if widths:
                set_cell_width(cells[i], widths[i])
            set_cell_text(cells[i], value, align=WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER)
    set_full_grid_borders(table)
    doc.add_paragraph()
    return table


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    add_footer(sec)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Title"]:
        if style_name in styles:
            styles[style_name].font.color.rgb = RGBColor(0, 0, 0)
    return doc


def add_cover(doc):
    top = doc.add_paragraph()
    top.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = top.add_run("密级★保密期限：公开                    合同编号：")
    set_run_font(r, east_asia="仿宋", ascii_font="Times New Roman", size=12)

    for text, size, bold, before in [
        ("专项科研课题", 20, True, 72),
        ("算法(模型)数据设计说明书", 22, True, 24),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        set_run_font(r, east_asia="黑体", ascii_font="Times New Roman", size=size, bold=bold)

    fields = [
        ("所属专题", "生产排产智能调度"),
        ("项目组长单位", ""),
        ("课题名称", "一号线日/周排产系统"),
        ("承研单位", ""),
        ("课题负责人", ""),
        ("联系电话", ""),
        ("研究周期", ""),
        ("编制日期", "2026年6月9日"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(3.0)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{label}：{value}")
        set_run_font(r, east_asia="仿宋", ascii_font="Times New Roman", size=14)
    doc.add_page_break()


def build_doc():
    doc = setup_doc()
    add_cover(doc)

    add_heading(doc, "1 概述", 1)
    add_heading(doc, "1.1 编写目的", 2)
    add_para(doc, "本文档依据《算法(模型)数据设计说明书》模板编写，用于统一一号线日/周排产系统全域数据口径，规范原始输入数据、约束数据、算法特征数据、输出结果数据的结构定义、流转规则、质量标准、存储方式与安全权限。文档明确数据全生命周期管控要求，避免数据口径不一致、脏数据干扰排产结果、接口字段不匹配等问题，为算法研发、功能测试、系统对接、交付运维和版本迭代提供统一的数据依据。")
    add_heading(doc, "1.2 文档边界范围", 2)
    add_para(doc, "本文档覆盖生产计划数据、型号工艺数据、约束参数数据、值班表数据、设备台数数据、故障窗口数据、动态重排数据、HGNN+PPO 参数数据、算法输出事件数据、算法对比指标数据，以及这些数据在前端 React 状态、后端 FastAPI/Pydantic 接口和算法求解过程中的流转规则。当前系统未接入数据库，本文同时给出可扩展的版本管理、存储目录和持久化建议。")
    add_heading(doc, "1.3 参考文档", 2)
    add_list(doc, [
        "日周排产系统功能说明.md",
        "需求修改调研报告.md",
        "排产系统意见评估.md",
        "算法设计文档",
        "src/App.jsx",
        "server.py",
        "scheduler_hgnn.py",
        "fjsp_hgnn_ppo.py",
        "行业通用数据治理规范",
    ])

    add_heading(doc, "2 数据总体架构设计", 1)
    add_heading(doc, "2.1 全链路数据流转流程", 2)
    add_para(doc, "系统数据流转流程为：前端默认配置/页面录入 → 原始业务数据校验 → 约束参数归一化 → 工序和设备实例展开 → 启发式排产或 HGNN+PPO 求解 → 业务约束映射与结果修正 → 排产事件输出 → 甘特图、工单表、算法对比报表和 PDF 导出。动态重排流程在已有排产事件基础上增加 NOW、故障窗口和插单任务，先完成冻结/作废分类，再对待排工序重新计算。")
    add_code_block(doc, "plan/types/cst/failures/hgnnParams\n  → validate material and duty roster\n  → normalize eqCount and dutyRosterByDay\n  → build timelines or FJSP proc\n  → runSchedule or /api/schedule/hgnn-ppo\n  → ScheduleEvent[]\n  → Gantt / WorkOrder / Metrics / PDF")
    add_heading(doc, "2.2 数据来源明细", 2)
    add_para(doc, "业务实时采集数据：前端页面实时录入的本周计划、型号配置、约束参数、值班表、设备台数、故障窗口和插单任务。")
    add_para(doc, "历史存量数据：当前代码未接入历史库，后续可将已生成计划、重排记录、导出工单和典型测试用例作为历史样本沉淀。")
    add_para(doc, "人工标注数据：本项目不属于监督分类/检测任务，没有传统人工标注标签；人工校准主要体现为计划员确认排产结果、故障原因和重排是否采纳。")
    add_para(doc, "公开补充数据集：当前不依赖公开数据集；HGNN+PPO 使用动态输入的 FJSP 问题实例在线求解。")
    add_heading(doc, "2.3 数据版本管理规范", 2)
    add_para(doc, "版本命名规则：建议采用 DATA_年月日_迭代序号，例如 DATA_20260609_V1。当前系统未落库时，可按导出的计划快照、约束快照和排产结果文件进行版本管理。")
    add_para(doc, "版本迭代触发条件：型号工艺字段调整、默认工时变化、值班表结构变化、设备台数规则变化、HGNN 参数口径变化、排产结果指标口径变化、动态重排规则变化。")
    add_para(doc, "版本留存规则：所有历史计划输入、约束快照和排产结果建议完整留存，禁止覆盖旧版本，以便模型复测、问题追溯和算法对比复现。")

    add_heading(doc, "3 数据集设计", 1)
    add_heading(doc, "3.1 划分原则", 2)
    add_para(doc, "当前系统没有固定离线训练集、验证集、测试集文件。HGNN+PPO 根据用户当前输入动态构造 FJSP 实例并在线训练/求解。为保证测试客观性，项目测试集应按业务场景划分，而非按随机样本划分，覆盖默认计划、不同型号组合、多批次、设备多台、人员空档、库存不足、设备故障和插单重排等场景。")
    add_para(doc, "划分比例：若后续沉淀历史排产样本，可按时间或计划批次划分，建议训练集 70%、验证集 15%、测试集 15%。对时序排产数据，应严格按照时间切割，避免未来计划或未来故障信息泄露到训练样本。")
    add_heading(doc, "3.2 防数据泄露关键规则", 2)
    add_list(doc, [
        "测试计划快照不得参与 HGNN 参数调优和默认规则选择。",
        "排产结果指标口径必须在训练、验证、测试和线上展示中保持一致。",
        "动态重排测试中，rescheduleAt 之后才发生的故障或插单不得提前写入初排输入。",
        "若后续使用历史数据训练模型，归一化、统计特征和参数调优只能基于训练集计算。",
        "测试集中的异常计划、故障窗口和人员空档应保持原始状态，不做人工修补后再评估。",
    ])
    add_heading(doc, "3.3 样本均衡策略", 2)
    add_para(doc, "本项目的样本均衡重点不是正负样本均衡，而是业务场景覆盖均衡。测试样本应覆盖 A/B/C 型号、短工时与长工时工序、单台和多台设备、不同人员容量、正常初排和动态重排等类别。对于罕见故障、极端库存不足、班次边界临界值等边缘场景，应单独建立回归测试样本，禁止只使用默认两批次计划作为唯一验证依据。")
    add_para(doc, "数据量级：当前演示规模以 3 个默认型号、5 类设备、5 个工作日、若干任务批次为主。后续建议至少沉淀 30 组以上典型计划快照、10 组以上动态重排案例、5 组以上异常约束案例，用于回归测试和版本验收。")

    add_heading(doc, "4 数据字段详细设计", 1)
    add_heading(doc, "4.1 原始接入数据字段定义", 2)
    add_grid_table(doc, ["字段名", "英文标识", "数据类型", "字段含义", "取值范围", "是否必填", "来源"], [
        ["任务标识", "plan.id", "string", "生产任务唯一标识", "如 t1、t2", "是", "本周计划页面"],
        ["型号标识", "plan.typeId", "string", "任务引用的产品型号", "ProductType.id", "是", "本周计划页面"],
        ["批次数", "plan.batches", "number", "任务生产批次数/锅数", "正整数", "是", "本周计划页面"],
        ["优先级", "plan.priority", "number", "排产优先顺序", "正整数，越小越优先", "是", "本周计划页面"],
        ["型号编码", "types.code", "string", "产品型号名称", "A型/B型/C型或自定义", "是", "型号配置页面"],
        ["工序名称", "ops.name", "string", "生产工序名称", "称量、预混、混合等", "是", "型号配置页面"],
        ["设备类型", "ops.eq", "string", "工序绑定设备类型", "称量台、搅拌机、混合锅、成型台、整装区", "是", "型号配置页面"],
        ["加工时长", "ops.dur", "number", "标准加工时长，单位小时", ">=0", "是", "型号配置页面"],
        ["清洗时长", "ops.cleanDur", "number", "工序后清洗占用时长", ">=0", "否", "型号配置页面"],
        ["AGV时长", "ops.agv", "number", "搬运占用时长", ">=0", "否", "型号配置页面"],
        ["所需人员", "ops.workers", "number", "工序执行所需人数", ">=1", "是", "型号配置页面"],
        ["库存", "cst.stockMatA/B/Release", "number", "原料与脱模剂可用库存", ">=0", "是", "约束参数页面"],
        ["班次时间", "cst.shiftStart/shiftEnd", "number", "每日生产班次起止小时", "0-24 且 start<end", "是", "约束参数页面"],
        ["值班人员", "dutyRosterByDay.names", "string[]", "指定工作日时段在岗人员姓名", "非空姓名数组", "是", "约束参数页面"],
        ["设备台数", "cst.eqCount", "object", "各设备类型可用台数", "1-10", "是", "约束参数页面"],
    ], widths=[2.5, 3.5, 2.2, 4.0, 4.0, 1.7, 3.0])
    add_heading(doc, "4.2 算法特征字段定义", 2)
    add_grid_table(doc, ["特征名称", "特征含义", "计算生成方式", "数据类型", "算法用途"], [
        ["工序可调度标记", "当前工序是否为某作业下一道待排工序", "由 FJSPEnv.next_op 判断", "float", "HGNN 工序节点特征"],
        ["工序完成标记", "工序是否已被调度", "检查 (job, op) 是否在 schedule 中", "float", "HGNN 工序节点特征"],
        ["作业最早开始时间", "当前批次下一工序可开始的最早时间", "job_time / MAX_T 归一化", "float", "HGNN 工序节点特征"],
        ["机器加工时间向量", "该工序在各机器上的加工/占用时长", "dur + cleanDur + agv，不可用机器置 0", "float[]", "HGNN 机器资格特征"],
        ["机器可用时间", "机器当前最早空闲时刻", "mach_time / MAX_T 归一化", "float", "HGNN 机器节点特征"],
        ["设备时间线", "设备实例已占用区间集合", "由已排事件和冻结事件生成", "list", "启发式冲突检查"],
        ["人员时间线", "各时段已占用人员数", "由已排事件 workers 字段累加", "list", "人员容量校验"],
        ["值班容量检查点", "工序区间内需要检查人员容量的时刻", "由人员占用边界和值班时段边界合并", "number[]", "跨时段人员约束"],
        ["故障窗口", "设备不可用时间段", "用户在动态重排弹窗录入", "Failure[]", "排产与重排避让"],
    ], widths=[3.2, 4.2, 5.0, 2.2, 4.4])
    add_heading(doc, "4.3 标签字段定义", 2)
    add_para(doc, "标签含义：本项目不是监督分类或检测任务，没有传统人工标注标签。用于算法评价的“标签”主要是业务验收标签和结果状态标签，包括排产事件状态、冻结状态、是否新排、是否完成、是否进行中、是否与故障冲突等。")
    add_para(doc, "标签编码规则：ScheduleEvent.status 使用 DONE、RUNNING 等字符串标识冻结状态；locked 使用 boolean 表示是否冻结；isNew 使用 boolean 表示是否为重排生成的新事件；算法对比中的较优项由指标方向自动判断。")
    add_para(doc, "正负样本定义：合法排产结果可视为正样本，违反设备重叠、人员超配、班次跨越、故障冲突、库存不足仍排产等规则的结果视为负样本或异常样本。")
    add_para(doc, "无效标签处理：状态缺失、工单号缺失、start/end 非法、opIdx 与型号工序不匹配的事件不得进入验收统计，应隔离并回溯生成逻辑。")

    add_heading(doc, "5 数据质量管控与清洗规则", 1)
    add_heading(doc, "5.1 缺失值处理规则", 2)
    add_para(doc, "缺失判定：plan、types、cst 中关键字段为空、类型错误、非法字符或引用不存在均视为缺失或无效。typeId 找不到对应型号、工序设备为空、值班时段 names 为空、班次时间缺失等会影响排产合法性。")
    add_para(doc, "低缺失率处理：对于兼容字段 dutyRoster、totalWorkers、workers，系统可通过 normalizeCst、normalizeDutySeg 自动补齐或归一化；对于 note、reason 等备注字段允许为空。")
    add_para(doc, "高缺失率处理：若计划无任务、型号无工序、约束无有效值班表或设备台数无法归一化，应阻止排产或恢复默认配置，避免错误数据进入算法。")
    add_heading(doc, "5.2 异常值处理规则", 2)
    add_para(doc, "异常判定：工序时长为负数、班次 start>=end、值班时段重叠、值班时段超出班次范围、设备台数小于 1、故障 end<=start、库存为负数、排产事件 start/end 非法等均为异常。")
    add_para(doc, "逻辑异常：同一设备实例事件时间重叠、同一批次工序顺序倒置、人员容量超出值班人数、工序跨越班次结束、故障窗口内安排新工序等属于逻辑异常。")
    add_para(doc, "处理方式：配置异常在页面侧提示并拦截；调度过程异常通过 findSlot 推进时间或返回 Infinity；后端求解失败时返回错误信息，前端可切换启发式算法。")
    add_heading(doc, "5.3 重复数据处理规则", 2)
    add_para(doc, "去重依据：任务以 id 去重，型号以 id 去重，工序以型号内 id 或 opIdx 去重，排产事件以 wo + opIdx 去重。动态重排合并 frozenMarked 与新排事件时，若已有相同 wo 和 opIdx 的冻结事件，应保留冻结事件并过滤新事件。")
    add_para(doc, "保留规则：对于已完成和正常进行中的工序，重排时优先保留原事件；对于被故障中断的事件，视为 cancelled 并从当前工序重新排产。")
    add_heading(doc, "5.4 数据一致性要求", 2)
    add_para(doc, "全链路统一时间单位为小时，以周一 00:00 为绝对时间零点；设备名称在前后端均使用同一中文设备类型和实例命名规则；物料字段统一为 matA、matB、release；HGNN 参数前端 epsClip、entropyCoef 在请求时映射为后端 eps_clip、entropy_coef；排产结果统一输出 ScheduleEvent 结构。")

    add_heading(doc, "6 标准化预处理流程", 1)
    add_para(doc, "步骤1：读取 plan、types、cst、failures、hgnnParams 等原始输入，统一对象结构和字段顺序。")
    add_para(doc, "步骤2：执行物料库存校验、值班表合法性校验、设备台数归一化和旧字段兼容转换。")
    add_para(doc, "步骤3：将 dutyRosterByDay 按工作日整理，names 归一化为 workers，空档作为不可排产时段处理。")
    add_para(doc, "步骤4：将 eqCount 展开为设备实例，单台保持原名，多台使用“设备名1、设备名2”的形式。")
    add_para(doc, "步骤5：计算每道工序总占用时长 dur + cleanDur + agv，并为 HGNN 构建 proc[job][op] 机器加工时间映射。")
    add_para(doc, "步骤6：初排时从零构建设备和人员时间线；重排时用 frozenEvents 初始化时间线，并用 batchProgress 跳过已完成工序。")
    add_para(doc, "步骤7：输出标准化 ScheduleEvent 数组，生成指标统计、甘特图和工单表。")

    add_heading(doc, "7 数据标注规范", 1)
    add_heading(doc, "7.1 标签体系规范", 2)
    add_para(doc, "系统标签体系围绕排产事件状态与验收结果设计。事件级标签包括 locked、isNew、status、isCleaning；重排统计标签包括 frozenCount、cancelledCount、newCount；指标级标签包括完工时长较优、利用率较优、清洗时长较优等。所有标签含义应在前端展示、接口返回和测试用例中保持一致。")
    add_heading(doc, "7.2 标注质检要求", 2)
    add_para(doc, "抽检比例：每次算法版本更新后，应抽检典型计划、异常约束和动态重排案例，建议不少于核心测试用例的 20%。")
    add_para(doc, "不合格标准：若出现设备时间重叠、人员超配、工序顺序错误、故障窗口冲突、重排丢失已完成工单、指标计算口径不一致等问题，应判定该批次测试不合格并返工修复。")
    add_para(doc, "歧义样本：对于业务上难以判断的排产结果，例如清洗是否单独展示、同类型多台设备故障影响范围、插单优先级冲突等，应由算法负责人和业务负责人共同确认口径。")

    add_heading(doc, "8 数据存储与目录规范", 1)
    add_heading(doc, "8.1 存储格式", 2)
    add_para(doc, "结构化业务配置：建议使用 JSON 保存 plan、types、cst、failures、hgnnParams 等对象快照。")
    add_para(doc, "模型训练数据集：当前系统无固定离线训练文件；若后续沉淀 FJSP 样本，可使用 JSON/NPY 保存 proc、机器数、最优结果和指标。")
    add_para(doc, "接口传输数据：前后端统一使用 JSON 格式，通过 FastAPI/Pydantic 模型进行字段校验。")
    add_para(doc, "结果归档数据：排产事件可保存为 JSON 或结构化表，工单可导出 PDF；若后续增加数据库，建议保存 schedule_plan、schedule_task、operation_def、constraint_snapshot、schedule_event、failure_window、reschedule_history 等表。")
    add_code_block(doc, "data/\n  snapshots/DATA_20260609_V1/\n    plan.json\n    types.json\n    constraints.json\n    events.json\n    failures.json\n    metrics.json\n  reports/\n    work_order.pdf\n    compare_report.json")

    add_heading(doc, "9 数据风险及应对方案", 1)
    add_para(doc, "数据量不足风险：当前样本主要来自默认配置和人工构造案例，覆盖面有限。应持续沉淀不同型号、批次、班次、设备台数和故障场景的计划快照。")
    add_para(doc, "样本不均衡风险：默认场景过多会导致测试只覆盖正常排产，忽略故障、空档、库存不足等边界。应建立异常场景清单并纳入回归测试。")
    add_para(doc, "数据漂移风险：后续工艺路线、设备数量、班次制度或人员值班方式变化时，旧数据口径可能不再适用。应通过数据版本号记录字段和规则变更。")
    add_para(doc, "数据质量差风险：非法时段、错误工时、缺少型号引用、设备名称不一致会直接影响排产合法性。应在前端输入、后端接口和算法执行前分别设置校验与归一化。")
    add_para(doc, "数据安全风险：值班人员姓名、生产计划和设备故障属于内部业务数据。部署到生产环境时，应增加访问控制、导出权限、操作审计和数据备份。")

    add_heading(doc, "10 文档变更记录", 1)
    add_grid_table(doc, ["版本号", "修改内容", "修改人", "修改日期"], [
        ["V1.0", "根据通用模板完成一号线日/周排产系统数据设计说明书初版。", "", "2026年6月9日"],
    ], widths=[3.0, 8.0, 3.0, 4.0])

    add_heading(doc, "11 附件", 1)
    add_list(doc, [
        "数据全链路流转流程图：可根据第 2.1 节流程绘制。",
        "字段样例数据表：可使用第 4 章字段表和第 8 章 JSON 快照示例生成。",
        "数据集目录结构截图：可按第 8.1 节建议目录创建后截图。",
        "数据质量检测脚本：可基于值班表校验、库存校验、事件合法性校验规则实现。",
    ])

    doc.save(OUT)
    repair_settings_zoom(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build_doc()
