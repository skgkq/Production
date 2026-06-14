from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from zipfile import ZipFile, ZIP_DEFLATED
import tempfile
import shutil


OUT = Path("一号线日周排产系统_算法设计文档_已填写.docx")


def set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=False, color="000000"):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_run_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        inserted = False
        for later_tag in ["w:tblLayout", "w:tblCellMar", "w:tblLook"]:
            later = tbl_pr.find(qn(later_tag))
            if later is not None:
                tbl_pr.insert(tbl_pr.index(later), borders)
                inserted = True
                break
        if not inserted:
            tbl_pr.append(borders)
    specs = {
        "top": ("single", "12"),
        "start": ("nil", "0"),
        "bottom": ("single", "12"),
        "end": ("nil", "0"),
        "insideH": ("nil", "0"),
        "insideV": ("nil", "0"),
    }
    for edge, (val, size) in specs.items():
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), val)
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:color"), "000000")


def set_full_grid_borders(table, color="D9DDE3", size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        inserted = False
        for later_tag in ["w:tblLayout", "w:tblCellMar", "w:tblLook"]:
            later = tbl_pr.find(qn(later_tag))
            if later is not None:
                tbl_pr.insert(tbl_pr.index(later), borders)
                inserted = True
                break
        if not inserted:
            tbl_pr.append(borders)
    specs = {
        "top": ("single", size),
        "start": ("single", size),
        "bottom": ("single", size),
        "end": ("single", size),
        "insideH": ("single", size),
        "insideV": ("single", size),
    }
    for edge, (val, sz) in specs.items():
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), val)
        elem.set(qn("w:sz"), sz)
        elem.set(qn("w:color"), color)


def set_cover_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=16, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.insert(0, tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))


def add_document_info_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("文档信息")
    set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size=24, bold=True)

    rows = [
        ["项目", "内容", "备注"],
        ["文档名称", "算法设计文档", "正式归档版"],
        ["项目名称", "PRODUCTION PLANNING · 一号线", ""],
        ["文档版本", "V1.0", "初始版本"],
        ["编写人", "", ""],
        ["审核人", "", ""],
        ["编制日期", "2026-06-08", ""],
        ["适用范围", "本项目算法研发、模型训练、代码实现、测试验收、交付运维全流程，为开发、测试、对接、迭代提供统一技术依据", ""],
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [5.0, 7.5, 5.0]
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            set_cell_width(cell, widths[c_idx])
            set_cover_cell_text(cell, value, bold=(r_idx == 0 or c_idx == 0))
            if r_idx == 7:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(15)
    set_full_grid_borders(table)
    doc.add_page_break()



def repair_settings_zoom(path):
    tmp_dir = Path(tempfile.mkdtemp(prefix="docx_repair_"))
    tmp_path = tmp_dir / path.name
    try:
        with ZipFile(path, "r") as zin, ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/settings.xml":
                    text = data.decode("utf-8")
                    text = text.replace('<w:zoom w:val="bestFit"/>', '<w:zoom w:val="bestFit" w:percent="100"/>')
                    data = text.encode("utf-8")
                zout.writestr(item, data)
        shutil.move(str(tmp_path), str(path))
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def add_table(doc, title, headers, rows):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = None
    r = cap.add_run(title)
    set_run_font(r, size=10.5, bold=True)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, align=WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER)
    set_table_borders(table)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level):
    p = doc.add_heading("", level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(6 if level == 1 else 3)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 3)
    run = p.add_run(text)
    set_run_font(run, east_asia="黑体", ascii_font="Times New Roman",
                 size={1: 16, 2: 14, 3: 12}.get(level, 12), bold=True)


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.first_line_indent = None
        run = p.add_run(item)
        set_run_font(run, size=12)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.right_indent = Pt(18)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, east_asia="等线", ascii_font="Consolas", size=10)
    run.font.color.rgb = RGBColor(31, 41, 55)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def build_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    add_footer(sec)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Title"]:
        if style_name in styles:
            styles[style_name].font.color.rgb = RGBColor(0, 0, 0)

    add_document_info_page(doc)

    add_heading(doc, "1 概述", 1)
    add_heading(doc, "1.1 编写目的", 2)
    add_para(doc, "本文档用于说明一号线日/周排产系统的算法设计方案，覆盖业务目标、输入输出、约束建模、启发式排产、HGNN+PPO 智能调度、动态重排、接口设计和测试验证要点，为后续开发、联调、汇报和维护提供统一依据。")
    add_heading(doc, "1.2 系统定位", 2)
    add_para(doc, "系统面向制造业一号线短周期生产计划，支持多型号、多批次、多工序的自动排产。核心任务是在班次、设备独占、同类型设备台数、值班人员容量、物料库存、清洗和 AGV 搬运等约束下，为工单生成可执行的工序时间轴，并在插单或设备故障发生后进行滚动重排。")
    add_heading(doc, "1.3 设计范围", 2)
    add_list(doc, [
        "覆盖前端启发式排产与后端 HGNN+PPO 排产两条算法路径。",
        "覆盖排产前物料校验、工作日班次和值班表校验、设备故障窗口避让、动态重排冻结与工单号稳定。",
        "不覆盖 ERP/MRP 级 BOM 展开、替代料、人员技能熟练度、成本会计和模具工装主数据管理。",
    ])

    add_heading(doc, "2 业务需求与约束", 1)
    add_heading(doc, "2.1 业务输入", 2)
    add_table(doc, "表 2-1 主要业务输入", ["输入项", "字段/来源", "说明"], [
        ["生产计划", "plan", "型号、批次数、优先级、备注；优先级数值越小越先排。"],
        ["型号工艺", "types.ops", "每个型号包含称量、预混、混合、模具装配、成型、整装等线性工序。"],
        ["资源约束", "cst", "库存、班次、工作日、值班表、设备台数。"],
        ["故障/插单", "failures/newTasks", "动态重排阶段输入，用于冻结已执行工序并重算后续工序。"],
    ])
    add_heading(doc, "2.2 资源约束", 2)
    add_table(doc, "表 2-2 资源与约束说明", ["约束", "设计说明"], [
        ["工序前后关系", "同一批次的工序按型号配置顺序串行执行，后继工序不得早于前序工序完成。"],
        ["设备独占", "同一设备实例同一时刻只能执行一道工序；当设备台数大于 1 时展开为同类型多个实例。"],
        ["班次边界", "工序不得跨越当日班次结束，剩余班次时间不足时顺延至下一工作日班次开始。"],
        ["值班表人员", "每个时段的在岗人数由姓名列表统计，工序执行期间任一检查点的占用人数不得超过容量。"],
        ["清洗/AGV", "清洗时长和 AGV 搬运时长计入工序总占用时长，其中清洗期间设备仍不可用。"],
        ["物料库存", "排产前按批次数汇总原料 A、原料 B、脱模剂消耗，不足则阻止生成排产。"],
        ["故障窗口", "故障窗口内设备不可用，排产和重排均需避让该时间段。"],
    ])
    add_heading(doc, "2.3 默认工艺路线", 2)
    add_table(doc, "表 2-3 默认型号工艺摘要", ["型号", "默认工序", "典型差异"], [
        ["A型", "称量 → 预混 → 混合 → 模具装配 → 成型 → 整装", "混合 3h、成型 4h。"],
        ["B型", "称量 → 预混 → 混合 → 模具装配 → 成型 → 整装", "预混 1.5h、混合 4h、成型 5h。"],
        ["C型", "称量 → 预混 → 混合 → 模具装配 → 成型 → 整装", "混合 2.5h、成型 3h，整体工时较短。"],
    ])

    add_heading(doc, "3 总体算法架构", 1)
    add_para(doc, "系统采用前后端协同的双算法架构。前端 React 应用内置启发式排产与动态重排逻辑，能够完整处理班次、值班表、设备台数和故障窗口；后端 FastAPI 提供 HGNN+PPO 智能调度接口，先对柔性作业车间问题进行策略求解，再映射回业务时间轴并进行实际约束修正。")
    add_code_block(doc, "计划/型号/约束 → 物料与值班表校验 → 选择算法\n  ├─ 启发式：按优先级展开批次 → 最早可行槽搜索 → 甘特图/工单\n  └─ HGNN+PPO：构建 FJSP → PPO 求解 → 约束映射 → 甘特图/工单\n动态重排：当前排程 + NOW + 故障/插单 → 冻结/作废分类 → 后续工序重排")
    add_heading(doc, "3.1 模块划分", 2)
    add_table(doc, "表 3-1 算法相关模块", ["模块", "文件", "职责"], [
        ["前端调度核心", "src/App.jsx", "实现 runSchedule、findSlot、findBestSlot、reschedule、指标计算和 UI 交互。"],
        ["后端 API", "server.py", "定义请求/响应模型，构建 FJSP 输入，调用 HGNN+PPO，并映射结果。"],
        ["HGNN+PPO 求解器", "scheduler_hgnn.py", "定义 FJSP 环境、异构图网络、Actor-Critic 和 PPO 训练循环。"],
        ["演示求解器", "fjsp_hgnn_ppo.py", "固定规模 FJSP 的原型与可视化演示代码。"],
    ])

    add_heading(doc, "4 数据模型设计", 1)
    add_heading(doc, "4.1 生产任务与工序", 2)
    add_para(doc, "生产任务由 plan 数组描述，每条任务包含 id、typeId、batches、priority 和 note。算法按 priority 升序展开为具体批次，并分配 WO-001 形式的工单号。型号 types 中的 ops 数组描述线性工序链，工序字段包括设备类型 eq、加工时长 dur、清洗 cleanDur、AGV agv、人员 workers 和物料消耗。")
    add_heading(doc, "4.2 设备实例", 2)
    add_para(doc, "设备基类固定为称量台、搅拌机、混合锅、成型台、整装区。约束参数 eqCount 用于描述同类型设备台数；当某类型台数为 1 时实例名保持原名，当台数大于 1 时展开为如“搅拌机1、搅拌机2”的设备实例。排产时同一工序仍绑定设备类型，算法在该类型的全部实例中选择最早可行槽。")
    add_heading(doc, "4.3 值班表", 2)
    add_para(doc, "值班表以 dutyRosterByDay 按工作日配置，每天包含多个时段，每个时段用 names 记录在岗人员姓名，workers 由姓名数量归一化得到。调度算法不区分个人技能，仅使用时段容量约束，满足当前系统“按姓名统计人数”的设计边界。")

    add_heading(doc, "5 启发式排产算法设计", 1)
    add_heading(doc, "5.1 算法思想", 2)
    add_para(doc, "启发式排产采用优先级驱动的最早可行槽搜索。算法先按任务优先级展开批次，再按每批工序顺序逐道调度。对每道工序，算法在该工序设备类型的可用实例中，寻找满足班次边界、设备空闲、故障避让和人员容量约束的最早开始时刻。")
    add_heading(doc, "5.2 核心流程", 2)
    add_list(doc, [
        "归一化约束参数，展开设备实例并初始化设备时间线。",
        "按优先级展开任务批次，为每批次生成稳定工单号。",
        "对批次内每道工序计算总占用时长：dur + cleanDur + agv。",
        "调用 findBestSlot 在同类型设备实例中选择最早可行槽。",
        "写入工序事件，更新设备时间线与人员占用时间线。",
        "按开始时间排序输出排产事件，用于甘特图、工单和指标统计。",
    ])
    add_heading(doc, "5.3 最早可行槽判定", 2)
    add_para(doc, "findSlot 从 minStart 开始迭代推进时间。若时刻不在工作班次内，则跳到下一可工作时刻；若工序跨越班次结束，则顺延到下一工作日；若与设备故障窗口或设备占用块重叠，则推进至冲突结束；若人员容量不足，则推进至下一个可能释放或值班时段边界。所有约束满足后返回该开始时刻。")
    add_code_block(doc, "start = nextWorkStart(minStart)\nwhile not feasible:\n    check shift boundary\n    check failure window\n    check equipment timeline\n    check duty-roster worker capacity\nreturn earliest feasible start")

    add_heading(doc, "6 HGNN+PPO 智能调度算法设计", 1)
    add_heading(doc, "6.1 FJSP 建模", 2)
    add_para(doc, "后端将生产批次映射为柔性作业车间调度问题。每个批次对应一个 job，每道生产工序对应一个 operation，同类型多台设备对应多个可选 machine。processing_time 采用加工、清洗、AGV 的合计时长。对于动态重排场景，已完成或允许继续执行的工序被作为 frozen_marked 注入时间线，待排部分从 batch_progress 指定的工序下标开始建模。")
    add_heading(doc, "6.2 异构图特征", 2)
    add_table(doc, "表 6-1 HGNN 图结构", ["关系", "矩阵/特征", "说明"], [
        ["工序前序关系", "A_prec", "同一 job 内 O(j,o) 到 O(j,o+1) 的顺序约束。"],
        ["机器资格关系", "A_elig", "表示某机器是否可加工某工序，用于机器与工序双向消息传递。"],
        ["机器关联关系", "A_conj", "机器节点全连接，表达机器间负载状态交互。"],
        ["工序特征", "op_f", "包含是否完成、是否可调度、作业最早开始时间、各机器加工时长。"],
        ["机器特征", "mach_f", "包含机器当前最早可用时间。"],
    ])
    add_heading(doc, "6.3 模型与训练", 2)
    add_para(doc, "HGNNR 使用关系特定图卷积分别处理前序关系、机器到工序、工序到机器和机器到机器的信息传递，再通过多头注意力融合工序视图。Actor-Critic 网络将工序嵌入与机器嵌入拼接，对每个可行动作 (job, operation, machine) 打分并采样；Critic 使用全局工序嵌入估计状态价值。训练采用 PPO，终局奖励为负 makespan，中间步奖励为 0。")
    add_table(doc, "表 6-2 默认 HGNN 参数", ["参数", "默认值", "说明"], [
        ["episodes", "300", "PPO 训练轮数，前端支持高级配置。"],
        ["lr", "0.0005", "Adam 学习率。"],
        ["gamma", "0.99", "折扣因子。"],
        ["eps_clip", "0.2", "PPO 裁剪范围。"],
        ["entropy_coef", "0.02", "熵正则系数，鼓励探索。"],
        ["d", "64", "图神经网络隐藏维度。"],
    ])
    add_heading(doc, "6.4 业务约束映射", 2)
    add_para(doc, "HGNN 原始环境只直接表达工序前后关系、机器可选集合和机器/作业可用时间。后端得到策略排程后，会按策略顺序调用业务 slot finder，将结果映射回真实班次、值班表、设备故障和冻结工序约束下的业务时间轴。因此，HGNN 负责提供较优的工序-设备选择与排序倾向，最终合法性由业务约束映射层保证。")

    add_heading(doc, "7 动态重排设计", 1)
    add_heading(doc, "7.1 触发条件", 2)
    add_para(doc, "动态重排由用户在排产结果页触发，输入重排基准时刻 NOW、设备故障窗口和可选插单任务。系统支持启发式与 HGNN+PPO 两条重排路径。")
    add_heading(doc, "7.2 冻结与作废规则", 2)
    add_table(doc, "表 7-1 动态重排分类规则", ["工序状态", "处理方式"], [
        ["已完成", "加入 frozen，保持原计划并标记 DONE，后续工序从下一道开始。"],
        ["进行中且未被故障中断", "加入 frozen，允许执行至原结束时刻并标记 RUNNING。"],
        ["进行中且被故障中断", "加入 cancelled，从当前工序重新排产。"],
        ["未开始", "不冻结，随后续待排工序一起重新计算。"],
        ["插单任务", "追加到重排计划中，按用户输入优先级参与后续排产。"],
    ])
    add_heading(doc, "7.3 工单号稳定", 2)
    add_para(doc, "重排时系统通过 assignStableWoForReschedule 保留已有任务的起始工单号，并为插单或新增批次分配未使用的新工单号。合并结果时以 wo 和 opIdx 识别已冻结工序，避免重复输出同一道工序。")

    add_heading(doc, "8 接口设计", 1)
    add_table(doc, "表 8-1 后端 API", ["接口", "方法", "输入", "输出"], [
        ["/api/schedule/hgnn-ppo", "POST", "ScheduleRequest(plan, types, cst, hgnn)", "events, makespan, hgnnParams"],
        ["/api/schedule/hgnn-ppo/reschedule", "POST", "RescheduleRequest(plan, types, cst, currentEvents, rescheduleAt, failures, hgnn)", "events, makespan, stats"],
    ])
    add_para(doc, "前端通过 buildHgnnRequestBody 将高级参数转换为后端字段，其中 epsClip 映射为 eps_clip，entropyCoef 映射为 entropy_coef。启发式路径在前端本地执行，不依赖后端服务。")

    add_heading(doc, "9 指标与报表", 1)
    add_para(doc, "系统提供算法对比报表，用同一计划分别执行启发式与 HGNN+PPO，并计算完工时长、设备平均利用率、产能利用率、清洗总时长、人工占用总时和工序总数。产能利用率口径为设备忙碌总时除以设备台数、工作日与日班次时长的乘积。")
    add_table(doc, "表 9-1 KPI 指标", ["指标", "计算口径", "优化倾向"], [
        ["完工时长", "所有事件 end 的最大值", "越低越好"],
        ["设备平均利用率", "各活跃设备 busy / makespan 的平均百分比", "越高越好"],
        ["产能利用率", "总忙碌时长 / 日历产能", "越高越好"],
        ["清洗总时长", "从 opName 中解析清洗时长并累加", "越低越好"],
        ["人工占用总时", "Σ 工序时长 × 所需人数", "越低越好"],
    ])

    add_heading(doc, "10 异常处理与边界", 1)
    add_list(doc, [
        "物料不足时前端弹窗提示并阻止排产。",
        "值班表存在重叠、无效时段、超出班次或无人时段配置错误时，启发式排产前拦截；空档作为不可排产时段处理。",
        "若排产超过工作日范围，nextWorkStart 返回 Infinity，界面以不可排或空结果提示。",
        "HGNN 服务不可用或返回异常时，前端提示智能调度失败，用户可切换启发式算法。",
        "系统仅按人数容量约束，不建模个人技能、个人熟练度和人工单价。",
    ])

    add_heading(doc, "11 测试验证建议", 1)
    add_table(doc, "表 11-1 验证用例", ["类别", "验证点", "预期结果"], [
        ["基础排产", "A型+B型各 1 批，默认 5 天班次", "输出完整工单，工序顺序正确，无设备重叠。"],
        ["物料校验", "库存低于计划消耗", "阻止排产并提示缺少物料。"],
        ["人员约束", "中午时段仅 1 人，安排 2 人工序", "工序避让该时段或顺延。"],
        ["设备台数", "搅拌机 eqCount=2", "甘特图展示搅拌机1/2，同类型工序可并行。"],
        ["故障重排", "NOW 后设备故障覆盖未完成工序", "冻结已完成/正常进行工序，故障冲突工序重新排。"],
        ["双算法对比", "同一计划运行启发式与 HGNN", "报表生成两套 KPI，可采用任一方案。"],
    ])

    add_heading(doc, "12 风险与改进方向", 1)
    add_para(doc, "当前 HGNN 奖励函数以 makespan 为单目标，尚未直接优化清洗成本、人工成本或负载均衡；相关指标更多用于报表对比。若后续要支持多目标权重，需要补充成本模型、奖励函数和训练数据。BOM 层级、替代料、模具绑定和人员熟练度属于更大范围的制造资源管理能力，不建议纳入当前一号线日/周排产核心算法范围。")
    add_para(doc, "后续可渐进增强的方向包括：设备故障精确到设备实例、HGNN 训练结果缓存、对比报表导出、按历史计划自动推荐 episodes，以及将设备利用率口径在报表、甘特图和导出 PDF 中统一。")

    doc.save(OUT)
    repair_settings_zoom(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT.resolve())
