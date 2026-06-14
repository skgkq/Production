from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from generate_algorithm_design_doc import (
    add_footer,
    add_heading,
    add_para,
    add_list,
    add_table,
    add_code_block,
    repair_settings_zoom,
    set_run_font,
)


OUT = Path("一号线日周排产系统_数据设计说明书_已填写.docx")


def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    add_footer(sec)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    return doc


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("一号线日/周排产系统数据设计说明书")
    set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size=18, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, value in [
        ("系统名称", "PRODUCTION PLANNING · 一号线"),
        ("文档版本", "v1.0"),
        ("依据版本", "项目代码 v4.1"),
        ("生成日期", "2026-06-08"),
        ("数据形态", "前端内存状态 + 后端 JSON 接口模型"),
    ]:
        r = subtitle.add_run(f"{label}：{value}\n")
        set_run_font(r, size=12)
    doc.add_page_break()


def build_doc():
    doc = setup_document()
    add_cover(doc)

    add_heading(doc, "1 概述", 1)
    add_heading(doc, "1.1 编写目的", 2)
    add_para(doc, "本文档用于说明一号线日/周排产系统的数据设计，包括核心业务对象、前端状态模型、后端接口模型、字段字典、数据流转、校验规则、数据安全和后续持久化扩展建议。文档依据项目中的 src/App.jsx、server.py、scheduler_hgnn.py 以及现有功能说明整理。")
    add_heading(doc, "1.2 数据设计范围", 2)
    add_list(doc, [
        "覆盖生产计划、型号工艺、约束参数、值班表、设备台数、排产结果、故障窗口、重排历史、HGNN 参数与算法对比报表。",
        "覆盖前端 React 状态对象与后端 FastAPI/Pydantic 请求响应模型之间的数据映射。",
        "当前系统未接入数据库，数据主要存在于运行时内存和接口 JSON 中；本文给出可落地的持久化表设计建议，但不把建议表描述为已实现数据库。",
        "不覆盖 ERP/MRP 级 BOM、替代料、人员技能档案、成本核算、模具工装主数据等未纳入当前项目的实体。",
    ])
    add_heading(doc, "1.3 术语说明", 2)
    add_table(doc, "表 1-1 术语说明", ["术语", "含义"], [
        ["型号", "产品工艺配置对象，如 A型、B型、C型。"],
        ["任务", "本周计划中的生产需求，包含型号、批次数和优先级。"],
        ["批次", "任务按 batches 展开的具体生产批，系统分配 WO-001 形式工单号。"],
        ["工序", "型号下的线性生产步骤，如称量、预混、混合、成型。"],
        ["设备类型", "工序绑定的资源类型，如搅拌机、混合锅。"],
        ["设备实例", "设备类型按 eqCount 展开的可排产实例，如搅拌机1、搅拌机2。"],
        ["值班表", "按工作日和时段配置在岗人员姓名，算法以姓名数量作为人员容量。"],
        ["排产事件", "一道工序在某设备实例上的计划开始、结束和业务属性。"],
    ])

    add_heading(doc, "2 数据总体架构", 1)
    add_para(doc, "系统采用轻量级前后端数据架构。前端以 React useState 保存业务配置、排产结果和交互状态；启发式算法在前端直接读取这些状态并输出事件列表。HGNN+PPO 智能调度通过 JSON 请求发送 plan、types、cst、hgnn 参数到 FastAPI 后端，后端基于 Pydantic 模型校验输入，调用求解器后返回 events 和 makespan。")
    add_code_block(doc, "前端状态\n  types / plan / cst / events / failures / reschedHistory / hgnnParams\n      │\n      ├─ 本地启发式 runSchedule → events\n      │\n      └─ POST /api/schedule/hgnn-ppo\n             plan + types + cst + hgnn → Pydantic 校验 → FJSP 求解 → events")
    add_heading(doc, "2.1 数据对象分层", 2)
    add_table(doc, "表 2-1 数据对象分层", ["层级", "对象", "说明"], [
        ["基础配置层", "ProductType、Operation、Constraints、DutySegment、EqCount", "描述型号工艺、库存、班次、人员和设备台数。"],
        ["计划输入层", "Task、Batch", "用户录入本周计划，算法展开为批次。Batch 为运行时派生对象。"],
        ["调度过程层", "EquipmentTimeline、WorkerTimeline、Failure", "算法内部用于判断设备占用、人员容量和故障避让。"],
        ["结果输出层", "ScheduleEvent、ScheduleMetrics、CompareReport", "用于甘特图、工单表、PDF 导出和算法对比。"],
        ["动态重排层", "FrozenEvent、BatchProgress、RescheduleHistory", "用于重排时保持已完成/进行中工序、工单号和历史记录。"],
        ["智能调度层", "HgnnParams、FJSP proc、HGNN schedule", "后端求解器使用的训练参数和中间结构。"],
    ])
    add_heading(doc, "2.2 数据存储现状", 2)
    add_para(doc, "当前代码未实现数据库、localStorage 或服务端持久化。页面刷新后，types、plan、cst、events 等运行时状态会恢复为默认值。后端服务仅在请求生命周期内构建 FJSP 输入和返回调度结果，不保存历史数据。")
    add_table(doc, "表 2-2 数据存储现状", ["数据类别", "当前存放位置", "生命周期"], [
        ["型号配置", "前端 useState(types)", "页面会话内有效。"],
        ["本周计划", "前端 useState(plan)", "页面会话内有效。"],
        ["约束参数", "前端 useState(cst)", "页面会话内有效。"],
        ["排产结果", "前端 useState(events)", "生成后保存在当前页面状态。"],
        ["故障与重排历史", "前端 useState(failures/reschedHistory)", "当前页面状态，刷新丢失。"],
        ["HGNN 求解输入输出", "HTTP 请求/响应内存对象", "单次请求内有效。"],
    ])

    add_heading(doc, "3 逻辑数据模型", 1)
    add_para(doc, "系统核心逻辑模型围绕“型号工艺 + 生产计划 + 约束参数 → 排产事件”展开。型号定义工序模板，任务引用型号并给出批次数，约束参数限定资源容量和可用时段，算法输出工序级事件。动态重排在已有事件上叠加故障窗口和插单任务，并通过冻结状态保持执行连续性。")
    add_code_block(doc, "ProductType 1 ── n Operation\nProductType 1 ── n Task\nTask 1 ── n Batch(派生)\nBatch 1 ── n ScheduleEvent\nConstraints 1 ── n DutySegment / EqCount\nFailure 0..n ── affects ScheduleEvent\nRescheduleHistory 记录一次重排的 NOW、故障数量和统计结果")
    add_heading(doc, "3.1 核心实体清单", 2)
    add_table(doc, "表 3-1 核心实体清单", ["实体", "实现位置", "主键/标识", "说明"], [
        ["ProductType", "src/App.jsx DEFAULT_TYPES / server.py TypeDef", "id", "产品型号及其工艺路线。"],
        ["Operation", "ProductType.ops / server.py OpDef", "id", "型号内的工序定义。"],
        ["Task", "plan / server.py TaskDef", "id", "本周计划中的生产任务。"],
        ["Constraints", "cst / server.py Constraints", "无单独 id", "库存、班次、工作日、值班表、设备台数。"],
        ["DutySegment", "dutyRosterByDay / DutySegDef", "id", "某工作日内的在岗时段。"],
        ["Failure", "failures / FailureDef", "运行时 id 或 eq+start+end", "设备故障窗口。"],
        ["ScheduleEvent", "events / EventDef", "wo + opIdx", "排产结果的最小业务单元。"],
        ["HgnnParams", "hgnnParams / HgnnParams", "无单独 id", "HGNN+PPO 求解参数。"],
    ])

    add_heading(doc, "4 数据字典", 1)
    add_heading(doc, "4.1 ProductType 型号数据", 2)
    add_table(doc, "表 4-1 ProductType 字段", ["字段", "类型", "必填", "说明"], [
        ["id", "string", "是", "型号内部唯一标识，如 pt1。"],
        ["code", "string", "是", "型号编码/名称，如 A型、B型。"],
        ["color", "string", "是", "前端展示颜色，用于甘特图和标签。"],
        ["ops", "Operation[]", "是", "该型号的线性工序数组，按数组顺序执行。"],
    ])
    add_heading(doc, "4.2 Operation 工序数据", 2)
    add_table(doc, "表 4-2 Operation 字段", ["字段", "类型", "默认值", "说明"], [
        ["id", "string", "无", "工序标识，同一型号内唯一。"],
        ["name", "string", "无", "工序名称，如称量、预混、混合。"],
        ["eq", "string", "无", "设备类型，取值来自称量台、搅拌机、混合锅、成型台、整装区。"],
        ["dur", "number", "无", "标准加工时长，单位小时。"],
        ["workers", "number", "1", "工序所需操作人数。"],
        ["cleanDur", "number", "0", "清洗时长，单位小时，计入设备占用。"],
        ["agv", "number", "0", "AGV 搬运时长，单位小时，计入工序总占用。"],
        ["isMix", "boolean", "false", "是否为混合工序，用于展示和工艺识别。"],
        ["matA", "number", "0", "原料 A 消耗量。"],
        ["matB", "number", "0", "原料 B 消耗量。"],
        ["release", "number", "0", "脱模剂消耗量。"],
    ])
    add_heading(doc, "4.3 Task 生产任务数据", 2)
    add_table(doc, "表 4-3 Task 字段", ["字段", "类型", "约束", "说明"], [
        ["id", "string", "必填唯一", "任务标识，如 t1。"],
        ["typeId", "string", "引用 ProductType.id", "计划生产的型号。"],
        ["batches", "number", "正整数", "本次任务的锅数/批次数。"],
        ["priority", "number", "正整数", "优先级，数值越小越先排。"],
        ["note", "string", "可为空", "备注字段，随工单输出。"],
        ["_woStart", "number", "运行时可选", "重排时用于保持工单号稳定的内部字段。"],
        ["_isNew", "boolean", "运行时可选", "插单任务标记。"],
    ])
    add_heading(doc, "4.4 Constraints 约束参数", 2)
    add_table(doc, "表 4-4 Constraints 字段", ["字段", "类型", "默认值", "说明"], [
        ["stockMatA", "number", "200", "原料 A 可用库存。"],
        ["stockMatB", "number", "80", "原料 B 可用库存。"],
        ["stockRelease", "number", "10", "脱模剂可用库存。"],
        ["shiftStart", "number", "8", "每日班次开始小时。"],
        ["shiftEnd", "number", "18", "每日班次结束小时。"],
        ["workDays", "number", "5", "排产覆盖工作日数量。"],
        ["dutyRosterByDay", "DutySegment[][]", "5 天默认值班表", "按工作日配置的值班表。"],
        ["dutyRoster", "DutySegment[]", "兼容字段", "旧版单日值班表，归一化后复制到工作日。"],
        ["eqCount", "Record<string, number>", "各设备 1 台", "设备类型到台数的映射，范围 1 到 10。"],
        ["totalWorkers", "number", "4", "兼容字段，无 dutyRoster 时生成默认人员。"],
        ["lunchStart/lunchEnd", "number", "12/13", "后端兼容字段，当前主要由值班表表达空档。"],
    ])
    add_heading(doc, "4.5 DutySegment 值班时段", 2)
    add_table(doc, "表 4-5 DutySegment 字段", ["字段", "类型", "约束", "说明"], [
        ["id", "string", "前端生成", "值班时段行标识。"],
        ["start", "number", "shiftStart <= start < end", "时段开始小时。"],
        ["end", "number", "start < end <= shiftEnd", "时段结束小时。"],
        ["names", "string[]", "至少 1 人", "在岗人员姓名列表。"],
        ["workers", "number", "由 names.length 归一化", "兼容字段，用于算法容量判断。"],
    ])
    add_heading(doc, "4.6 Failure 故障窗口", 2)
    add_table(doc, "表 4-6 Failure 字段", ["字段", "类型", "说明"], [
        ["eq", "string", "故障设备实例名称，如搅拌机或搅拌机1。"],
        ["start", "number", "故障开始绝对小时，以周一 00:00 为 0。"],
        ["end", "number", "故障结束绝对小时，必须大于 start。"],
        ["reason", "string", "故障原因，可为空。"],
    ])
    add_heading(doc, "4.7 ScheduleEvent 排产事件", 2)
    add_table(doc, "表 4-7 ScheduleEvent 字段", ["字段", "类型", "说明"], [
        ["wo", "string", "工单号，如 WO-001。"],
        ["taskId", "string", "来源任务 id，重排和工单号稳定使用。"],
        ["batchLabel", "string", "展示标签，如 A型-批1。"],
        ["batchNum", "number", "任务内批次序号。"],
        ["ptId/ptCode/ptColor", "string", "型号标识、编码和展示颜色。"],
        ["opName", "string", "工序名称，包含清洗/AGV 附加说明。"],
        ["opIdx", "number", "型号工序数组下标。"],
        ["eq", "string", "实际排入的设备实例。"],
        ["start/end", "number", "计划开始/结束绝对小时。"],
        ["dur", "number", "总占用时长，单位小时。"],
        ["workers", "number", "所需人员数。"],
        ["isCleaning", "boolean", "当前实现中工序与清洗合并，通常为 false。"],
        ["note", "string", "任务备注。"],
        ["locked", "boolean", "动态重排中冻结事件标记。"],
        ["isNew", "boolean", "是否为重排后新生成事件。"],
        ["status", "string", "DONE/RUNNING 等冻结状态。"],
    ])

    add_heading(doc, "5 接口数据设计", 1)
    add_heading(doc, "5.1 HGNN 排产请求", 2)
    add_table(doc, "表 5-1 ScheduleRequest", ["字段", "类型", "说明"], [
        ["plan", "TaskDef[]", "生产计划任务数组。"],
        ["types", "TypeDef[]", "型号和工艺配置数组。"],
        ["cst", "Constraints", "约束参数。"],
        ["episodes", "number", "兼容字段，默认 300。"],
        ["hgnn", "HgnnParams | null", "完整 HGNN 参数配置。"],
    ])
    add_heading(doc, "5.2 HGNN 动态重排请求", 2)
    add_table(doc, "表 5-2 RescheduleRequest", ["字段", "类型", "说明"], [
        ["plan", "TaskDef[]", "原计划和可能新增的插单任务。"],
        ["types", "TypeDef[]", "当前型号工艺配置。"],
        ["cst", "Constraints", "当前约束参数。"],
        ["currentEvents", "EventDef[]", "重排前已有排产事件。"],
        ["rescheduleAt", "number", "重排基准时刻 NOW。"],
        ["failures", "FailureDef[]", "本次参与重排的故障窗口。"],
        ["episodes/hgnn", "number/HgnnParams", "智能调度参数。"],
    ])
    add_heading(doc, "5.3 响应数据", 2)
    add_table(doc, "表 5-3 响应结构", ["接口", "主要字段", "说明"], [
        ["/api/schedule/hgnn-ppo", "events, makespan, hgnnParams", "返回智能调度后的工序事件、完工时长和实际参数。"],
        ["/api/schedule/hgnn-ppo/reschedule", "events, makespan, stats", "返回重排后事件、完工时长和冻结/新增统计。"],
        ["前端启发式", "events", "本地 runSchedule 直接返回事件数组。"],
        ["算法对比", "greedy.metrics, hgnn.metrics", "前端派生对比报表，不是后端固定响应模型。"],
    ])

    add_heading(doc, "6 数据流转设计", 1)
    add_heading(doc, "6.1 初始排产数据流", 2)
    add_list(doc, [
        "用户在本周计划、型号配置和约束参数中维护 plan、types、cst。",
        "生成排产前，前端汇总物料消耗并校验库存，校验值班表时段合法性。",
        "启发式模式直接调用 runSchedule；HGNN 模式调用后端接口。",
        "算法返回 events 后，前端更新排产结果页、甘特图、工单表和 KPI。",
    ])
    add_heading(doc, "6.2 动态重排数据流", 2)
    add_list(doc, [
        "用户输入 rescheduleAt、故障窗口和插单任务。",
        "系统根据 currentEvents 分类 frozen、cancelled 和待排工序。",
        "assignStableWoForReschedule 为已有任务保留工单号，为插单分配新工单号。",
        "重排算法合并 frozenMarked 与新排事件，输出新的 events。",
        "前端记录 reschedHistory，并在甘特图和工单中显示 locked、DONE、RUNNING、isNew 状态。",
    ])
    add_heading(doc, "6.3 算法对比数据流", 2)
    add_para(doc, "算法对比功能在同一份 plan、types、cst 上分别执行启发式和 HGNN+PPO，随后调用 computeScheduleMetrics 生成完工时长、设备平均利用率、产能利用率、清洗总时长、人工占用总时和工序总数。对比报表仅派生展示，不修改原始计划数据；用户点击采用方案后才把对应 events 写入当前排产结果。")

    add_heading(doc, "7 数据校验规则", 1)
    add_table(doc, "表 7-1 主要校验规则", ["对象", "规则", "处理方式"], [
        ["Task.batches", "应为正整数，前端步进器限制范围。", "不满足时无法形成有效批次。"],
        ["Task.typeId", "必须引用存在的 ProductType.id。", "调度展开时找不到型号则跳过。"],
        ["Operation.dur", "应为非负数，业务上加工工时建议大于 0。", "作为总时长组成部分参与排产。"],
        ["Constraints.eqCount", "每类设备台数归一化到 1 到 10。", "超出范围自动截断。"],
        ["DutySegment", "start < end，且在班次范围内，不得与上一时段重叠。", "页面红色提示，生成前拦截。"],
        ["DutySegment.names", "至少 1 人，workers 由姓名数量计算。", "无人时段视为不可用或配置错误。"],
        ["库存", "计划总消耗不得超过 stockMatA/stockMatB/stockRelease。", "不足时弹窗提示并阻止排产。"],
        ["Failure", "end 必须大于 start，并绑定具体设备实例。", "用于重排和故障窗口避让。"],
        ["ScheduleEvent", "同设备实例不应时间重叠，同批次 opIdx 应递增。", "由调度算法保证。"],
    ])

    add_heading(doc, "8 数据安全与权限", 1)
    add_para(doc, "当前系统为本地演示型排产应用，不包含用户登录、角色权限、数据库账号或外部商业数据接口。数据主要为工艺参数、值班姓名和生产计划，属于内部业务数据。若部署到生产环境，应增加身份认证、接口鉴权、操作审计和敏感字段脱敏。")
    add_table(doc, "表 8-1 数据安全建议", ["数据", "风险", "建议"], [
        ["值班人员姓名", "涉及人员信息。", "生产部署时应限制访问范围，导出 PDF 可按需要隐藏姓名。"],
        ["生产计划与工艺参数", "涉及产线计划和工艺能力。", "接口增加鉴权，导出文件控制分发范围。"],
        ["HGNN 参数与结果", "可能影响排产决策。", "保留默认参数，限制高级参数编辑权限。"],
        ["故障与重排历史", "反映设备运行状态。", "持久化后应记录操作者和时间，便于追溯。"],
    ])

    add_heading(doc, "9 持久化扩展建议", 1)
    add_para(doc, "若后续需要把系统从演示状态升级为可长期使用的排产工具，建议增加轻量数据库层。根据当前数据模型，可优先设计以下表结构；这些表为建议设计，不代表当前代码已经落库。")
    add_table(doc, "表 9-1 建议持久化表", ["表名", "主键", "主要字段", "说明"], [
        ["product_type", "id", "code, color, created_at, updated_at", "保存型号基本信息。"],
        ["operation_def", "id", "type_id, seq, name, eq, dur, workers, clean_dur, agv, mat_a, mat_b, release", "保存型号工序定义。"],
        ["schedule_plan", "id", "name, work_days, shift_start, shift_end, status", "保存一次排产计划头信息。"],
        ["schedule_task", "id", "plan_id, type_id, batches, priority, note", "保存计划任务明细。"],
        ["constraint_snapshot", "id", "plan_id, stock_mat_a, stock_mat_b, stock_release, eq_count_json", "保存排产时约束快照。"],
        ["duty_segment", "id", "plan_id, day_idx, start, end, names_json", "保存按日值班表。"],
        ["schedule_event", "id", "plan_id, wo, task_id, batch_num, op_idx, eq, start, end, dur, status", "保存排产结果事件。"],
        ["failure_window", "id", "plan_id, eq, start, end, reason", "保存故障窗口。"],
        ["reschedule_history", "id", "plan_id, reschedule_at, frozen_count, new_count, created_at", "保存重排历史。"],
    ])
    add_heading(doc, "9.1 JSON 字段建议", 2)
    add_para(doc, "对于 eqCount、names、hgnnParams、算法对比报表等结构较小且变化频繁的字段，可以采用 JSON 字段保存快照，避免过早拆分出过多配置表。对于 operation_def 和 schedule_event 这类需要查询、排序和追溯的核心业务数据，应使用结构化列。")

    add_heading(doc, "10 数据备份与恢复", 1)
    add_para(doc, "当前系统无持久化层，因此备份与恢复主要依赖项目代码和导出的 PDF/Word 文档。若增加数据库，建议以排产计划为备份粒度，保存 plan、types、cst、events 和 failures 的完整快照；恢复时先还原基础配置，再还原排产事件和重排历史。")
    add_table(doc, "表 10-1 备份对象建议", ["对象", "备份频率", "恢复用途"], [
        ["型号与工序配置", "每次修改后", "恢复工艺参数。"],
        ["计划任务", "每次生成排产前", "复现输入计划。"],
        ["约束快照", "每次生成排产时", "复现排产约束。"],
        ["排产事件", "每次生成或重排后", "恢复甘特图和工单。"],
        ["重排历史", "每次动态重排后", "追溯插单、故障和冻结规则。"],
    ])

    add_heading(doc, "11 与算法设计的关系", 1)
    add_para(doc, "数据设计直接服务于算法设计。启发式算法依赖 Operation、Constraints、DutySegment、Failure 和 ScheduleEvent 判断最早可行槽；HGNN+PPO 将 Task 和 Operation 转换为 FJSP 的 proc 结构，将 eqCount 展开为机器集合，并在结果映射阶段重新应用业务约束。动态重排通过 ScheduleEvent 的 locked、status、isNew 字段表达冻结、进行中和新排状态。")

    add_heading(doc, "12 附录：典型 JSON 示例", 1)
    add_code_block(doc, """{
  "plan": [{"id": "t1", "typeId": "pt1", "batches": 1, "priority": 1, "note": ""}],
  "types": [{
    "id": "pt1",
    "code": "A型",
    "color": "#3b82f6",
    "ops": [{"id": "a1", "name": "称量", "eq": "称量台", "dur": 0.5, "workers": 1}]
  }],
  "cst": {
    "stockMatA": 200,
    "shiftStart": 8,
    "shiftEnd": 18,
    "workDays": 5,
    "eqCount": {"称量台": 1, "搅拌机": 1, "混合锅": 1, "成型台": 1, "整装区": 1}
  }
}""")

    doc.save(OUT)
    repair_settings_zoom(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT.resolve())
