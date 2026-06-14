from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from generate_algorithm_design_doc import (
    add_code_block,
    add_document_info_page,
    add_footer,
    add_heading,
    add_list,
    add_para,
    add_table,
    repair_settings_zoom,
    set_cell_text,
    set_cell_width,
    set_full_grid_borders,
    set_run_font,
)


OUT = Path("一号线日周排产系统_算法设计文档_已填写.docx")


def add_grid_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths is None
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        if widths:
            set_cell_width(cell, widths[i])
        set_cell_text(cell, h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if widths:
                set_cell_width(cells[i], widths[i])
            set_cell_text(cells[i], value, align=WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER)
    set_full_grid_borders(table)
    doc.add_paragraph()
    return table


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    add_footer(sec)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Title"]:
        if style_name in styles:
            styles[style_name].font.color.rgb = RGBColor(0, 0, 0)
    return doc


def build_doc():
    doc = setup_doc()
    add_document_info_page(doc)

    add_heading(doc, "1 概述", 1)
    add_heading(doc, "1.1 编写目的", 2)
    add_para(doc, "本文档依据《算法设计文档（通用模板）》编写，用于明确一号线日/周排产系统的算法设计思路、核心原理、整体架构、流程逻辑、参数配置、性能指标、约束条件及落地规范。文档统一开发、测试、模型训练、系统对接和后续迭代的认知，保障排产算法可实现、可复现、可维护、可扩展。")
    add_heading(doc, "1.2 业务背景与问题描述", 2)
    add_para(doc, "业务场景：本系统服务于制造业一号线日/周生产计划编排，面向计划员、生产管理人员和排产执行人员，在多型号、多批次、多工序的生产条件下，根据班次、设备、人员、库存、清洗和 AGV 搬运等约束自动生成工序级排产结果，并提供甘特图、工单表、PDF 导出和动态重排能力。")
    add_para(doc, "现有痛点：人工排产需要同时考虑工序前后关系、设备独占、人员时段容量、清洗占用、物料库存和突发故障，规则多且联动强，容易出现设备冲突、人员超配、班次跨越、插单响应慢和工单号混乱等问题。传统静态表格方案难以快速评估不同算法结果，也不便在故障后保留已执行工序并滚动重排。")
    add_para(doc, "算法介入价值：系统通过启发式最早可行槽搜索保障快速生成合法排程，通过 HGNN+PPO 智能调度探索更优的工序-设备选择和排序倾向，通过动态重排机制处理插单与故障，实现生产计划从人工经验编排向自动化、可视化、可比较、可追溯的调度决策转变。")
    add_heading(doc, "1.3 设计目标与验收指标", 2)
    add_para(doc, "功能目标：实现生产计划录入、型号工艺配置、约束参数配置、启发式排产、HGNN+PPO 智能排产、双算法对比、动态重排、甘特图展示、工单表输出和 PDF 导出。算法需输出每道工序的工单号、批次、工序、设备实例、开始时间、完成时间、工时、人员和状态。")
    add_para(doc, "性能目标：默认计划规模下，启发式算法应在前端交互级时间内完成排产；HGNN+PPO 默认 episodes=300，用于智能调度和算法对比。系统应计算完工时长、设备平均利用率、产能利用率、清洗总时长、人工占用总时和工序总数等指标。")
    add_para(doc, "稳定性目标：算法应在缺少可行时段、人员空档、设备故障、库存不足、后端智能调度失败等异常情况下给出明确提示或降级路径；动态重排应保持已完成和正常进行中工序不被破坏，并保留已有工单号。")
    add_para(doc, "验收标准：排产结果中同一批次工序顺序正确，同一设备实例同一时刻不重叠，工序不跨越班次边界，人员占用不超过值班表容量，库存不足时阻止排产，故障窗口内不安排新工序，重排后冻结工序状态清晰且工单号稳定。")
    add_heading(doc, "1.4 适用边界与约束", 2)
    add_para(doc, "适用范围：适用于一号线固定工艺路线、日/周周期、工作日班次、多型号多批次、多设备类型及同类型多台设备的生产排程场景。默认设备类型包括称量台、搅拌机、混合锅、成型台、整装区。")
    add_para(doc, "不适用范围：当前算法不覆盖 ERP/MRP 级 BOM 展开、替代料决策、个人技能熟练度建模、工装/模具主数据绑定、完整制造成本最优和跨产线协同调度。")
    add_para(doc, "外部约束：前端启发式调度依赖浏览器运行环境；HGNN+PPO 后端依赖 Python、FastAPI、PyTorch、NumPy 等环境；当前系统数据主要为运行时状态，未接入数据库持久化。")

    add_heading(doc, "2 总体方案设计", 1)
    add_heading(doc, "2.1 整体设计思路", 2)
    add_para(doc, "整体方案采用“规则可行性保障 + 智能调度优化探索”的双算法路线。前端启发式算法以业务约束为核心，按优先级展开批次并逐工序寻找最早可行槽，保证排产结果在班次、设备、人员、故障和物料约束下可执行。后端 HGNN+PPO 将问题转化为柔性作业车间调度 FJSP，通过异构图神经网络提取工序、机器和前序关系特征，再使用 PPO 学习动作策略，用于探索更优完工时长和设备分配。")
    add_heading(doc, "2.2 算法整体架构", 2)
    add_para(doc, "算法架构分为数据接入层、数据校验层、特征与约束构建层、核心算法层、结果修正层和输出对接层。数据接入层读取本周计划、型号工艺、约束参数、故障窗口和插单任务；数据校验层完成库存和值班表合法性检查；核心算法层执行启发式或 HGNN+PPO；结果修正层处理冻结、故障避让、工单号稳定和指标计算。")
    add_code_block(doc, "数据接入 → 库存/值班表校验 → 设备实例展开/约束构建\n  ├─ 启发式最早可行槽搜索\n  └─ HGNN+PPO FJSP 求解 + 业务约束映射\n→ 排产事件合并 → KPI 计算 → 甘特图/工单/PDF/对比报表")
    add_heading(doc, "2.3 技术选型说明", 2)
    add_grid_table(doc, ["技术模块", "选型方案", "选型理由", "替代方案与劣势"], [
        ["核心算法模型", "启发式最早可行槽搜索 + HGNN-R + PPO", "启发式保障规则合法和快速响应；HGNN+PPO 支持柔性作业车间智能优化展示。", "纯人工/静态规则难以处理复杂重排；单一强化学习方案训练成本高且解释性弱。"],
        ["框架/工具", "React + Vite + FastAPI + PyTorch + NumPy", "前端适合交互式配置和甘特图展示；FastAPI 便于 JSON 接口；PyTorch 支持 HGNN 与 PPO 实现。", "传统后端模板页面交互弱；仅前端实现深度学习不可行；重型 MES 平台接入成本高。"],
        ["训练策略", "PPO on-policy 训练，终局奖励为 -makespan", "动作空间为可行的工序-机器组合，PPO 实现简单、稳定，适合演示型智能调度。", "DQN 难处理动态动作集合；监督学习缺少标注最优解。"],
        ["推理加速方案", "默认 episodes=300，可配置 episodes、隐藏维度和学习率", "在效果与耗时之间折中，前端提供高级参数入口。", "全量超参开放会增加误用风险；过大 episodes 影响交互体验。"],
    ], widths=[3.6, 4.3, 5.8, 5.8])
    add_heading(doc, "2.4 整体业务流程", 2)
    add_para(doc, "业务流程为：计划员录入本周生产任务，维护型号工艺参数和约束参数；系统在生成排产前校验物料库存和值班表；用户选择启发式或 HGNN+PPO 算法后生成排产事件；结果页展示甘特图、工单、设备利用率和完工时间；若发生插单或设备故障，用户输入重排基准时刻和故障窗口，系统冻结已完成/进行中的合法工序，作废被故障中断的工序，并对剩余工序和插单任务重新排产。")

    add_heading(doc, "3 核心算法详细设计", 1)
    add_heading(doc, "3.1 核心原理", 2)
    add_para(doc, "启发式算法核心原理是约束驱动的最早可行槽搜索。算法按任务优先级展开批次，并按型号工序顺序依次调度。每道工序的总占用时长为加工时长、清洗时长和 AGV 搬运时长之和。算法在该工序绑定设备类型的所有设备实例中，选择满足班次、设备占用、故障窗口和人员容量约束的最早开始时间。")
    add_para(doc, "HGNN+PPO 核心原理是将排产问题建模为柔性作业车间调度。每个批次视为一个作业，每道工序视为作业内的操作，同类型多台设备视为可选机器集合。HGNN-R 通过前序关系、机器资格关系和机器关联关系进行消息传递，Actor-Critic 网络对可行动作打分，PPO 根据完工时长奖励更新策略。")
    add_heading(doc, "3.2 算法输入输出定义", 2)
    add_grid_table(doc, ["类型", "字段/参数名称", "数据类型", "取值说明/含义", "是否必填"], [
        ["输入参数", "plan", "Task[]", "生产任务数组，包含型号、批次数、优先级、备注。", "是"],
        ["输入参数", "types", "ProductType[]", "型号工艺数组，包含工序、设备、工时、清洗、人员和物料。", "是"],
        ["输入参数", "cst", "Constraints", "库存、班次、工作日、值班表、设备台数等约束。", "是"],
        ["输入参数", "failures", "Failure[]", "动态重排时的设备故障窗口。", "否"],
        ["输入参数", "hgnn", "HgnnParams", "episodes、lr、gamma、eps_clip、entropy_coef、d。", "否"],
        ["输出结果", "events", "ScheduleEvent[]", "工序级排产结果，包含工单号、设备、开始结束时间和状态。", "是"],
        ["输出结果", "makespan", "number", "全部工序最大完工时刻，单位为绝对小时。", "是"],
        ["输出结果", "stats/metrics", "object", "重排统计或算法对比指标，如利用率、清洗时长、人工占用。", "否"],
    ], widths=[2.2, 3.8, 3.2, 7.8, 2.2])
    add_heading(doc, "3.3 核心计算流程与伪代码", 2)
    add_para(doc, "启发式排产按“批次展开 → 工序遍历 → 设备实例选择 → 约束校验 → 事件写入”执行。若某个候选时刻不满足约束，算法会推进到班次开始、故障结束、设备冲突结束、人员释放或值班时段边界等下一个可能可行点，直到找到可行槽或超出工作日范围。")
    add_code_block(doc, "for task in sort_by_priority(plan):\n  for batch in expand(task.batches):\n    prev_end = min_start\n    for op in product_type.ops:\n      dur = op.dur + op.cleanDur + op.agv\n      best = INF\n      for eq_instance in instances_for_type(op.eq):\n        t = find_slot(eq_instance, prev_end, dur, workers, failures)\n        best = min(best, t)\n      emit ScheduleEvent(best)\n      update equipment timeline and worker timeline\n      prev_end = best.end")
    add_para(doc, "动态重排先按 NOW 对现有事件分类：已完成工序冻结为 DONE；进行中且未被故障中断的工序冻结为 RUNNING；进行中且被故障窗口覆盖的工序作废，从当前工序重新排；未开始工序全部重新计算；插单任务追加到计划中，并使用稳定工单号规则分配新的 WO 编号。")
    add_heading(doc, "3.4 关键参数设计", 2)
    add_grid_table(doc, ["参数名称", "默认值", "可调范围", "参数作用", "调优影响"], [
        ["episodes", "300", "50-2000", "PPO 训练轮数。", "越大可能找到更优排程，但耗时增加。"],
        ["lr", "0.0005", "0.0001-0.01", "Adam 学习率。", "过大不稳定，过小收敛慢。"],
        ["gamma", "0.99", "0.8-0.999", "折扣因子。", "影响终局奖励传播。"],
        ["eps_clip", "0.2", "0.05-0.4", "PPO 策略更新裁剪范围。", "过大易震荡，过小更新慢。"],
        ["entropy_coef", "0.02", "0-0.2", "熵正则系数。", "提高探索，过大影响收敛。"],
        ["d", "64", "16-256", "HGNN 隐藏维度。", "维度越大表达能力越强，计算开销越高。"],
        ["eqCount", "各设备 1 台", "1-10", "同类型设备台数。", "台数越多并行能力越强，甘特图设备实例增加。"],
        ["shiftStart/shiftEnd", "8/18", "0-24", "每日班次范围。", "影响可排产时间窗口。"],
    ], widths=[3.0, 2.7, 3.0, 5.6, 5.2])
    add_heading(doc, "3.5 模型结构设计（深度学习类算法必填）", 2)
    add_para(doc, "网络整体结构：HGNN-R 包含工序编码器、机器编码器、四类关系卷积模块和跨关系多头注意力融合模块。关系卷积分别处理工序前序关系、机器到工序资格关系、工序到机器资格关系和机器间关联关系。Actor 将工序嵌入与机器嵌入拼接后输出动作得分，Critic 使用工序全局嵌入估计状态价值。")
    add_para(doc, "层级参数：工序输入维度为 3+n_machines，包含是否完成、是否可调度、作业最早开始时间和各机器加工时长；机器输入维度为 1，表示机器当前最早可用时间；隐藏维度默认 d=64；多头注意力默认 n_heads=4；激活函数采用 ReLU，归一化采用 LayerNorm。")
    add_para(doc, "损失函数设计：PPO 损失由策略裁剪损失、价值函数均方误差和熵正则组成。终局奖励为负 makespan，中间步骤奖励为 0，目标是最小化全部工序最大完工时间。")
    add_para(doc, "优化器与学习率策略：使用 Adam 优化器，默认学习率 5e-4；StepLR 按训练轮数分段衰减；梯度裁剪 max_norm=1.0，降低训练震荡风险。")
    add_heading(doc, "3.6 规则与后处理策略", 2)
    add_para(doc, "后处理策略包括：对 HGNN 原始排程按策略顺序重新调用业务可行槽搜索，确保结果满足班次、值班表、设备故障和冻结工序约束；对动态重排结果合并 frozenMarked 与新排事件，避免重复输出相同 wo+opIdx；对工序名称追加清洗和 AGV 说明；对 KPI 指标进行统一计算；对不可行或异常结果以提示、降级或空结果方式处理。")

    add_heading(doc, "4 数据方案设计", 1)
    add_heading(doc, "4.1 数据来源与数据构成", 2)
    add_para(doc, "算法数据来源于前端页面录入和系统默认配置，包括生产计划 plan、型号工艺 types、约束参数 cst、动态重排故障 failures、插单任务 newTasks 和智能调度参数 hgnnParams。当前系统未接入数据库，数据以 React 运行时状态和接口 JSON 形式存在。")
    add_table(doc, "表 4-1 数据构成", ["数据对象", "主要字段", "用途"], [
        ["Task", "id, typeId, batches, priority, note", "描述本周生产任务。"],
        ["ProductType", "id, code, color, ops", "描述型号和工艺路线。"],
        ["Operation", "name, eq, dur, workers, cleanDur, agv, matA, matB, release", "描述工序资源需求和物料消耗。"],
        ["Constraints", "stockMatA, shiftStart, dutyRosterByDay, eqCount", "描述库存、班次、人员和值班约束。"],
        ["ScheduleEvent", "wo, opIdx, eq, start, end, status", "描述算法输出的工序级计划。"],
    ])
    add_heading(doc, "4.2 数据集划分规则", 2)
    add_para(doc, "当前 HGNN+PPO 采用动态输入的在线训练/求解方式，没有固定离线训练集、验证集和测试集文件。测试验证以业务场景用例划分，包括默认计划、物料不足、人员空档、多设备台数、故障重排和算法对比等场景。若后续沉淀历史排产数据，可按时间划分训练、验证和测试集合，避免同一计划快照跨集合造成数据泄露。")
    add_heading(doc, "4.3 数据预处理规范", 2)
    add_para(doc, "前端和后端均提供约束归一化逻辑。eqCount 会补全默认设备并限制在 1 到 10；dutyRosterByDay 会补足 5 天并将 names 归一化为 workers；旧字段 dutyRoster 可复制到全周；工序总时长统一计算为 dur + cleanDur + agv；FJSP 建模时按设备台数展开机器索引，并按 batchProgress 跳过已完成工序。")
    add_heading(doc, "4.4 特征工程设计（根据模型实际情况）", 2)
    add_para(doc, "HGNN 特征主要包括工序节点特征和机器节点特征。工序节点特征由完成标记、当前可调度标记、作业最早开始时间和各机器加工时间组成；机器节点特征为机器当前最早可用时间。图结构特征包括同一作业内的前序关系、工序与可用机器之间的资格关系，以及机器间全连接关系。启发式算法则直接使用时间线特征，即设备占用时间线、人员占用时间线、值班容量检查点和故障窗口。")

    add_heading(doc, "5 训练与调优方案", 1)
    add_heading(doc, "5.1 训练环境配置", 2)
    add_para(doc, "HGNN+PPO 后端运行环境为 Python 服务，依赖 FastAPI、Uvicorn、Pydantic、PyTorch、NumPy。前端为 React 18 + Vite，负责配置输入、结果展示和算法对比。默认 CPU 即可运行演示规模问题；若扩大任务数量、设备数量或 episodes，可使用具备更高 CPU/GPU 资源的环境。")
    add_heading(doc, "5.2 超参数配置", 2)
    add_para(doc, "当前开放的主要超参数为 episodes、lr、gamma、eps_clip、entropy_coef 和隐藏维度 d。前端默认提供 episodes 简易配置和高级参数入口，默认参数为调优后的演示配置。生产/汇报场景建议使用默认参数或少量调整 episodes，避免非算法人员误调学习率、裁剪范围等敏感参数。")
    add_heading(doc, "5.3 训练流程", 2)
    add_list(doc, [
        "根据 plan、types、cst 展开批次和设备实例，构建 proc[job][op] = {machine_idx: processing_time}。",
        "初始化 FJSPEnv、HGNN-R、Actor-Critic、Adam 优化器和学习率调度器。",
        "每个 episode 重置环境，循环选择可行动作并记录轨迹。",
        "全部工序完成后获得 makespan，使用 reward=-makespan 计算折扣回报。",
        "按 PPO 裁剪目标更新策略网络和价值网络，保存最优 makespan 对应排程。",
        "将最优策略排程映射回业务排产事件，并执行班次、人员、故障等约束修正。",
    ])
    add_heading(doc, "5.4 调优策略", 2)
    add_para(doc, "调优方向包括：增加 episodes 提升搜索充分性；适当降低学习率提升稳定性；调整 entropy_coef 平衡探索和收敛；在规模扩大时提高隐藏维度 d。若出现结果波动，可固定随机种子、提高 episodes、增加算法对比次数或使用启发式结果作为业务兜底。若耗时过高，应降低 episodes、缩小设备和任务规模，或对同一计划结果做缓存。")

    add_heading(doc, "6 性能与效果评估", 1)
    add_heading(doc, "6.1 评估指标", 2)
    add_table(doc, "表 6-1 评估指标", ["指标", "计算方式", "目标方向"], [
        ["完工时长", "所有事件 end 的最大值", "越低越好"],
        ["设备平均利用率", "活跃设备 busy/makespan 的平均值", "越高越好"],
        ["产能利用率", "总忙碌时长 /（设备台数 × 工作日 × 日班次时长）", "越高越好"],
        ["清洗总时长", "从工序名称中解析清洗时长累加", "越低越好"],
        ["人工占用总时", "Σ 工序时长 × 所需人数", "越低越好"],
        ["合法性", "设备不重叠、人员不超配、班次不跨越、故障不冲突", "必须满足"],
    ])
    add_heading(doc, "6.2 测试数据集要求", 2)
    add_para(doc, "测试数据应覆盖默认 A/B/C 型号、不同批次数和优先级、单台/多台设备、不同值班表容量、库存不足、故障窗口、插单任务和 HGNN 参数变化。测试计划应保证输入数据独立可复现，并记录 plan、types、cst、failures 和 hgnnParams 快照。")
    add_heading(doc, "6.3 预期效果与基线对比", 2)
    add_para(doc, "基线方案为人工排产或前端启发式排产。HGNN+PPO 的预期效果是在部分计划上获得更短完工时长或更高设备利用率；启发式方案的预期效果是快速、稳定、可解释地生成合法排程。系统通过算法对比报表展示两种方案在完工时长、利用率、清洗总时长和人工占用上的差异，用户可选择采用任一方案。")
    add_para(doc, "异常和极端场景下，算法应优先保证输出合法性和可解释性。当后端智能调度不可用、训练耗时过长或无可行解时，系统可提示失败并切换启发式算法；当物料不足或值班表非法时，排产前直接拦截。")

    add_heading(doc, "7 接口与集成设计", 1)
    add_heading(doc, "7.1 算法服务接口定义", 2)
    add_table(doc, "表 7-1 算法服务接口", ["接口", "方法", "入参", "出参"], [
        ["/api/schedule/hgnn-ppo", "POST", "plan, types, cst, episodes, hgnn", "events, makespan, hgnnParams"],
        ["/api/schedule/hgnn-ppo/reschedule", "POST", "plan, types, cst, currentEvents, rescheduleAt, failures, hgnn", "events, makespan, stats"],
        ["前端 runSchedule", "本地函数", "plan, types, cst, opts", "events"],
        ["前端 reschedule", "本地函数", "currentEvents, plan, types, cst, rescheduleAt, failures, newTasks", "events, stats"],
    ])
    add_heading(doc, "7.2 系统集成方式", 2)
    add_para(doc, "启发式算法以内嵌方式集成在前端 src/App.jsx 中，适合离线或后端不可用时使用。HGNN+PPO 以服务化方式集成在 FastAPI 后端 server.py 中，前端通过 fetch 调用接口。两种算法共用统一的数据结构和输出事件模型，因此甘特图、工单表、PDF 导出、动态重排状态展示和算法对比报表无需区分底层算法。")
    add_heading(doc, "7.3 输入输出示例", 2)
    add_code_block(doc, "请求示例：\n{\n  \"plan\": [{\"id\":\"t1\",\"typeId\":\"pt1\",\"batches\":1,\"priority\":1,\"note\":\"\"}],\n  \"types\": [{\"id\":\"pt1\",\"code\":\"A型\",\"ops\":[{\"name\":\"称量\",\"eq\":\"称量台\",\"dur\":0.5,\"workers\":1}]}],\n  \"cst\": {\"shiftStart\":8,\"shiftEnd\":18,\"workDays\":5,\"eqCount\":{\"称量台\":1}},\n  \"hgnn\": {\"episodes\":300,\"lr\":0.0005,\"gamma\":0.99}\n}\n\n返回示例：\n{\n  \"events\": [{\"wo\":\"WO-001\",\"batchLabel\":\"A型-批1\",\"opName\":\"称量\",\"eq\":\"称量台\",\"start\":8,\"end\":8.5}],\n  \"makespan\": 8.5,\n  \"hgnnParams\": {\"episodes\":300}\n}")

    add_heading(doc, "8 容错与兼容性设计", 1)
    add_heading(doc, "8.1容错设计", 2)
    add_para(doc, "系统对物料不足、值班表非法、故障窗口冲突、后端智能调度失败、无可行排产时段等异常进行处理。物料不足和非法值班表在排产前拦截；故障窗口在 findSlot 中避让；动态重排通过 frozen/cancelled 分类保护已执行工序；HGNN 接口失败时前端提示错误并可切换启发式算法；重排合并时用 wo+opIdx 去重，防止重复工序。")
    add_heading(doc, "8.2  兼容性设计", 2)
    add_para(doc, "前端兼容现代浏览器运行环境，构建工具为 Vite；后端兼容 Python 3 环境，使用 FastAPI 和 PyTorch。数据模型保留 dutyRoster、totalWorkers、lunchStart、lunchEnd 等旧字段兼容逻辑，实际使用中优先采用 dutyRosterByDay 和 names。设备实例命名兼容单台设备原名与多台设备编号形式。")

    add_heading(doc, "9风险分析与应对措施", 1)
    add_para(doc, "数据风险：当前系统缺少真实历史排产数据和离线训练集，HGNN+PPO 主要基于动态输入在线求解。应沉淀典型计划快照、约束快照和重排案例，逐步建立测试样本库。")
    add_para(doc, "效果风险：PPO 终局奖励仅为 -makespan，对清洗成本、人工成本和负载均衡没有直接优化。应在界面明确 HGNN 当前为完工时长优化，并把清洗和人工作为报表指标或后续多目标奖励扩展。")
    add_para(doc, "性能风险：episodes 增大或任务规模扩大时，训练耗时可能影响交互体验。应限制默认参数、提供加载提示、支持启发式兜底，并考虑结果缓存。")
    add_para(doc, "落地风险：当前数据未持久化，刷新页面会丢失配置和结果。若进入实际生产使用，应增加数据库、权限、审计和备份恢复机制。")
    add_para(doc, "迭代风险：若后续引入 BOM、替代料、人员技能或模具工装约束，将显著扩展数据模型和算法复杂度，应作为独立版本规划，不宜直接叠加到当前核心排产逻辑中。")

    add_heading(doc, "10 附则", 1)
    add_heading(doc, "10.1 文档变更记录", 2)
    add_grid_table(doc, ["版本号", "修改内容", "修改人", "修改日期"], [
        ["V1.0", "根据通用模板完成一号线日/周排产系统算法设计文档初版。", "", "2026-06-08"],
    ], widths=[3.0, 8.0, 3.0, 4.0])
    add_heading(doc, "10.2 参考文档", 2)
    add_list(doc, [
        "日周排产系统功能说明.md",
        "需求修改调研报告.md",
        "排产系统意见评估.md",
        "src/App.jsx",
        "server.py",
        "scheduler_hgnn.py",
        "fjsp_hgnn_ppo.py",
        "数据设计说明书",
    ])
    add_heading(doc, "10.3 附件", 2)
    add_para(doc, "可附加算法流程图、接口请求样例、典型排产结果截图、算法对比报表、动态重排案例、核心代码片段和测试验收记录。当前文档已在正文中给出核心伪代码、接口示例和参数表。")

    doc.save(OUT)
    repair_settings_zoom(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build_doc()
